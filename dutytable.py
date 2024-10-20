import calendar
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from datetime import datetime

def get_default_date():
    current_date = datetime.now()
    return current_date.year, current_date.month

def get_user_input(default_year, default_month):
    year = input(f"Введите год [{default_year}]: ")
    month = input(f"Введите месяц (1-12) [{default_month}]: ")
    year = int(year) if year else default_year
    month = int(month) if month else default_month
    num_duty = int(input("Введите количество дежурных: "))
    return year, month, num_duty

def create_document():
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    return doc

def add_table_header(table, month_days):
    days_ru = ["Пн", "Вт", "Ср", "Чт", "Пт", "Сб", "Вс"]
    for i, (day, weekday) in enumerate(month_days):
        cell = table.cell(0, i + 1)
        cell.width = Inches(1.5)
        date_text = f'{days_ru[weekday]} {day}'
        p = cell.add_paragraph()
        run = p.add_run(date_text)
        run.font.size = Pt(8)  # Установить размер шрифта для дат
        if weekday in (5, 6):  # 5 - суббота, 6 - воскресенье
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет для выходных

def add_duty_names(table, num_duty):
    for i in range(1, num_duty + 1):
        cell = table.cell(i, 0)
        cell.width = Inches(2)
        cell.text = f'Дежурный {i}'

def format_table_cells(table, month_days):
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = Inches(1.5)
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)
            cell_paragraph = cell.paragraphs[0]
            cell_paragraph.paragraph_format.left_indent = Inches(0.1)
            cell_paragraph.paragraph_format.right_indent = Inches(0.1)
            cell_paragraph.paragraph_format.space_before = Pt(5)
            cell_paragraph.paragraph_format.space_after = Pt(5)
            # Закрасить столбцы выходных дней светло-красным
            if j > 0 and month_days[j-1][1] in (5, 6):  # 5 - суббота, 6 - воскресенье
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'FFCCCC')  # Светло-красный цвет
                tcPr.append(shading_elm)

def main():
    default_year, default_month = get_default_date()
    year, month, num_duty = get_user_input(default_year, default_month)

    months_ru = ["", "Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", 
                 "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь", "Декабрь"]

    cal = calendar.Calendar()
    month_days = [day for day in cal.itermonthdays2(year, month) if day[0] != 0]

    doc = create_document()
    doc.add_heading(f'График дежурств на {months_ru[month]} {year}', level=1)

    table = doc.add_table(rows=num_duty + 1, cols=len(month_days) + 1)

    add_table_header(table, month_days)
    add_duty_names(table, num_duty)
    format_table_cells(table, month_days)

    doc_name = f'График_дежурств_{month}_{year}.docx'
    doc.save(doc_name)
    print(f'Документ сохранен как {doc_name}')

if __name__ == "__main__":
    main()